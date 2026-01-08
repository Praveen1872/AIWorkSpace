import streamlit as st
from google import genai
from docx import Document
from io import BytesIO
import re
import firebase_admin
from firebase_admin import credentials, firestore

is_logged_in = st.session_state.get('logged_in', False)
if not is_logged_in:
    st.switch_page("pages/login.py")

def initialize_firebase():
    if not firebase_admin._apps:
        try:
            cred = credentials.Certificate({
                "type": "service_account",
                "project_id": os.getenv("FIREBASE_PROJECT_ID"),
                "private_key": os.getenv("FIREBASE_PRIVATE_KEY").replace("\\n", "\n"),
                "client_email": os.getenv("FIREBASE_CLIENT_EMAIL"),

                # 🔑 REQUIRED EXTRA FIELDS
                "token_uri": "https://oauth2.googleapis.com/token",
                "auth_uri": "https://accounts.google.com/o/oauth2/auth",
            })

            firebase_admin.initialize_app(cred)

        except Exception as e:
            raise RuntimeError(f"Firebase Initialization Failed: {e}")

    return firestore.client()


firestore_db = initialize_firebase()
user_uid = st.session_state.get("user_uid")
if "ai_suggestions" not in st.session_state:
    st.session_state.ai_suggestions = []

if "change_log" not in st.session_state:
    st.session_state.change_log = []

word_col = (
    firestore_db
    .collection("users")
    .document(user_uid)
    .collection("documents")
    .document("word")
    .collection("items")
)
def shorten_title(text, max_len=60):
    text = re.sub(r'\s+', ' ', text.strip())
    return text if len(text) <= max_len else text[:max_len].rsplit(" ", 1)[0] + "…"

h_cols = st.columns([2, 0.7, 0.7, 0.7, 0.9, 1.8, 1], vertical_alignment="center")
with h_cols[0]: 
    st.markdown("<h3 style='margin:0;'>🚀 AI Mentor</h3>", unsafe_allow_html=True)

with h_cols[1]: 
    if st.button("PPT 🖼️", use_container_width=True): st.switch_page("pages/ppt_editor.py")
with h_cols[2]: 
    if st.button("Word 📝", use_container_width=True,type="primary"): st.switch_page("pages/word_editor.py")
with h_cols[3]: 
    if st.button("Notes 📓", use_container_width=True): st.switch_page("pages/note.py")
with h_cols[4]: 
    if st.button("Summarizer 📝", use_container_width=True): st.switch_page("pages/Summarizer.py")
with h_cols[6]:
    if st.button("Logout 🚪", use_container_width=True):
        st.session_state.logged_in = False
        st.switch_page("AIMentor.py")

st.markdown("<hr style='margin:0 0 20px 0;'>", unsafe_allow_html=True)


def clean_text(text):
    # Remove markdown headings and bold
    text = re.sub(r'#+\s*', '', text)
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'\*', '', text)
    return text.strip()


from docx.shared import Pt, RGBColor

def create_docx(content):
    doc = Document()

    # Define styles
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)

    for line in content.split('\n'):
        line = line.strip()
        if not line:
            continue

        # ===== HEADINGS =====
        if line.isupper() or len(line) < 60 and line.endswith(':'):
            heading = doc.add_heading(clean_text(line), level=1)
            heading.style.font.size = Pt(16)
            heading.style.font.color.rgb = RGBColor(0, 0, 128)
            continue

        # ===== BULLETS =====
        if line.startswith('-') or line.startswith('*'):
            p = doc.add_paragraph(
                clean_text(line),
                style='List Bullet'
            )
            p.style.font.size = Pt(11)
            continue

        # ===== PARAGRAPH =====
        p = doc.add_paragraph(clean_text(line))
        p.paragraph_format.space_after = Pt(8)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

import os 
try:
    client = genai.Client(api_key=os.getenv("GEMINI_API_KEY"))
except Exception as e:
    st.error(f"Gemini Initialization Failed: {e}")
def store_word_chunks(word_doc_ref, full_text):
    chunks_col = word_doc_ref.collection("chunks")

    paragraphs = [
        p.strip() for p in full_text.split("\n")
        if len(p.strip()) > 40
    ]

    for p in paragraphs:
        chunks_col.add({
            "text": p,
            "embedding": []  # placeholder for future RAG
        })

st.title("📄 Report Assistant")
st.write("Refine and export your research into polished Microsoft Word documents.")

topic = st.text_input("Report Topic", placeholder="e.g., The impact of renewable energy on global economy")

if st.button("Generate Report", use_container_width=True):
    if topic:
        with st.spinner("AI Mentor is drafting your report..."):
            try:
                prompt = (
                    f"Write a comprehensive formal academic report about: {topic}. "
                    "Use ## for section headings."
                )

                response = client.models.generate_content(
                    model="gemini-2.5-flash-lite",
                    contents=[prompt],
                    config=genai.types.GenerateContentConfig(
                        system_instruction=(
                            "You are a professional academic mentor. "
                            "Write formal, well-structured reports. "
                            "Do not use markdown bolding (**) in titles or headers."
                        )
                    )
                )

                # ✅ STORE IN SESSION STATE
                st.session_state.report_text = response.text

                # ✅ CREATE FIRESTORE DOC
                word_doc_ref = word_col.document()
                word_doc_ref.set({
                    "title": shorten_title(topic),
                    "created_at": firestore.SERVER_TIMESTAMP,
                    "source": "word_generator"
                })
                st.session_state.active_word_id = word_doc_ref.id
                # ✅ ALWAYS READ FROM SESSION STATE
                store_word_chunks(
                    word_doc_ref,
                    st.session_state.report_text
                )

                st.session_state.active_word_id = word_doc_ref.id
                st.success("Draft completed!")

            except Exception as e:
                st.error(f"Generation failed: {e}")
    else:
        st.warning("Please enter a topic to begin.")
if "report_text" in st.session_state:
    final_text = st.text_area("Review & Edit Draft", value=st.session_state.report_text, height=500)
    word_file = create_docx(final_text)
    safe_filename = clean_text(topic).replace(' ', '_')[:20]
    
    st.download_button(
        label="📥 Download as Word (.docx)",
        data=word_file,
        file_name=f"Report_{safe_filename}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True
    )
firestore_db = firestore.client()
def load_user_words():
    docs = word_col.order_by(
        "created_at",
        direction=firestore.Query.DESCENDING
    ).stream()

    return [{"id": d.id, **d.to_dict()} for d in docs]
with st.sidebar:
    st.subheader("📂 Word History")

    words = load_user_words()

    if words:
        title_map = {
            f"{i+1}. {shorten_title(w['title'])}": w["id"]
            for i, w in enumerate(words)
        }

        selected = st.selectbox(
            "Past Documents",
            title_map.keys()
        )

        if st.button("📂 Load Document"):
            st.session_state.active_word_id = title_map[selected]
            st.rerun()

    if st.button("🗑️ Clear All Word Docs"):
        for doc in word_col.stream():
            for c in doc.reference.collection("chunks").stream():
                c.reference.delete()
            doc.reference.delete()

        st.session_state.pop("word_text", None)
        st.rerun()
def load_user_words():
    docs = word_col.order_by(
        "created_at",
        direction=firestore.Query.DESCENDING
    ).stream()

    return [{"id": d.id, **d.to_dict()} for d in docs]
if "active_word_id" in st.session_state:
    word_doc = word_col.document(st.session_state.active_word_id)
    chunks = word_doc.collection("chunks").stream()

    full_text = "\n\n".join(
        c.to_dict()["text"] for c in chunks
    )

    st.session_state.word_text = full_text
def extract_rewrite(ai_text):
    match = re.search(r"Suggested Rewrite:\s*(.*?)\nReason:", ai_text, re.S)
    return match.group(1).strip() if match else ai_text.strip()

st.markdown("### ✨ AI Assist (Edit Selected Text)")

selected_text = st.text_area(
    "Paste the text you want AI to edit or explain",
    height=150,
    placeholder="Paste a paragraph or bullet point here..."
)

action = st.selectbox(
    "Choose AI Action",
    [
        "Improve writing (clarity & flow)",
        "Make it more formal",
        "Simplify for beginners",
        "Fix grammar only",
        "Explain this (Doubt Solver)",
        "Summarize this"
    ]
)

if st.button("🤖 Apply AI", use_container_width=True):
    if not selected_text.strip():
        st.warning("Please paste some text first.")
    else:
        with st.spinner("AI is working..."):
            ai_prompt = f"""
You are a professional academic editor.

TASK: {action}

RULES:
- Edit ONLY the given text
- Do NOT add markdown (#, **)
- Keep meaning intact unless summarizing
- Output plain text only

TEXT:
{selected_text}
"""
            response = client.models.generate_content(
                model="gemini-2.5-flash-lite",
                contents=[ai_prompt]
            )

            st.session_state.ai_result = response.text
if "ai_result" in st.session_state:
    st.markdown("### ✅ AI Output")
    st.text_area(
        "AI Edited Text",
        value=st.session_state.ai_result,
        height=150
    )

    if st.button("🔁 Replace in Document"):
        if "report_text" in st.session_state:
            st.session_state.report_text = (
                st.session_state.report_text
                .replace(selected_text, st.session_state.ai_result)
            )
            st.success("Text replaced successfully!")
            st.rerun()
st.markdown("## 💬 AI Comment Suggestions")

comment_text = st.text_area(
    "Paste text to get AI suggestions (no auto edit)",
    height=120,
    placeholder="Paste a paragraph or sentence here..."
)

comment_action = st.selectbox(
    "AI Suggestion Type",
    [
        "Improve clarity",
        "Make more formal",
        "Fix grammar",
        "Simplify explanation",
        "Explain meaning (Doubt Solver)"
    ]
)

if st.button("💡 Get AI Suggestion", use_container_width=True):
    if not comment_text.strip():
        st.warning("Paste some text first.")
    else:
        with st.spinner("AI is analyzing..."):
            prompt = f"""
You are an academic writing assistant.

TASK: {comment_action}

RULES:
- Do NOT use markdown
- Do NOT change meaning unless simplifying
- Provide TWO parts only:

FORMAT:
Suggested Rewrite:
<text>

Reason:
<short explanation>

TEXT:
{comment_text}
"""
            response = client.models.generate_content(
                model="gemini-2.5-flash-lite",
                contents=[prompt]
            )

            st.session_state.ai_suggestions.append({
                "original": comment_text,
                "suggestion": response.text,
                "timestamp": firestore.SERVER_TIMESTAMP
            })
for i, sug in enumerate(st.session_state.ai_suggestions):
    st.markdown("---")
    st.markdown(f"### 💡 Suggestion {i+1}")

    st.markdown("**Original Text:**")
    st.code(sug["original"])

    st.markdown("**AI Suggestion:**")
    st.code(sug["suggestion"])

    col1, col2 = st.columns(2)

    with col1:
        if st.button("✅ Apply", key=f"apply_{i}"):
            st.session_state.report_text = (
                st.session_state.report_text
                .replace(sug["original"], extract_rewrite(sug["suggestion"]), 1)
            )

            st.session_state.change_log.append({
                "original": sug["original"],
                "new": extract_rewrite(sug["suggestion"]),
                "status": "applied"
            })

            st.success("Change applied")
            st.rerun()

    with col2:
        if st.button("❌ Ignore", key=f"ignore_{i}"):
            st.session_state.change_log.append({
                "original": sug["original"],
                "new": None,
                "status": "ignored"
            })
            st.success("Suggestion ignored")
            st.rerun()
st.markdown("## 🧾 Track Changes")

if not st.session_state.change_log:
    st.info("No changes yet.")
else:
    for c in st.session_state.change_log:
        st.markdown("---")
        st.markdown(f"**Status:** {c['status'].upper()}")
        st.markdown("**Original:**")
        st.code(c["original"])

        if c["new"]:
            st.markdown("**AI Applied Version:**")
            st.code(c["new"])
